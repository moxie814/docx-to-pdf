ECHO is on.


import os
from flask import Flask, request, render_template, send_from_directory, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from reportlab.pdfgen import canvas
from io import BytesIO

app = Flask(__name__)

# Folder for storing uploaded files and converted PDFs
UPLOAD_FOLDER = 'uploads'
CONVERTED_FOLDER = 'converted'
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_FOLDER'] = CONVERTED_FOLDER

# Create folders if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

# Check allowed extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Home route
@app.route('/')
def index():
    return render_template('index.html')

# Route to handle file upload
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Extract metadata
        doc = Document(filepath)
        metadata = {
            "author": doc.core_properties.author,
            "title": doc.core_properties.title,
            "subject": doc.core_properties.subject,
            "word_count": len(doc.paragraphs)
        }

        # Convert docx to PDF
        pdf_filename = f"{filename.rsplit('.', 1)[0]}.pdf"
        pdf_filepath = os.path.join(app.config['CONVERTED_FOLDER'], pdf_filename)
        convert_to_pdf(filepath, pdf_filepath)

        return render_template('index.html', metadata=metadata, pdf_filename=pdf_filename)
    return jsonify({"error": "Invalid file format"}), 400

# Convert docx to PDF
def convert_to_pdf(docx_filepath, pdf_filepath):
    doc = Document(docx_filepath)
    buffer = BytesIO()

    c = canvas.Canvas(buffer)
    y_position = 800
    for para in doc.paragraphs:
        c.drawString(50, y_position, para.text)
        y_position -= 20
        if y_position < 50:
            c.showPage()
            y_position = 800

    c.save()

    with open(pdf_filepath, 'wb') as f:
        f.write(buffer.getvalue())

# Route to download the PDF
@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['CONVERTED_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True)
